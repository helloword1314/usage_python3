# -*- coding: utf-8 -*-
# ---------------------------------------------------------------------------
# .py
# Created on: 2020-04-28 22:45:37.00000
# Description: 修改空间参考后，点位的值也会发生改变
# ---------------------------------------------------------------------------

# Import arcpy module

import os
import sys
import docx
from docx import Document


frame_path ='\\'.join(os.path.realpath(__file__).split('\\')[:-2])
module_python=frame_path+"\cwframe_pythion3\module_for_python"



sys.path.append(module_python)
import basic_python.basic_function_python as basic_function_python
import basic_python.file_handle_python as file_handle_python
import basic_python.file_stream_python as file_stream_python
import basic_python.file_docx_python as file_docx_python
 
arrDocx=["laoyouji_4_.docx","laoyouji_5_.docx","laoyouji_6_.docx"]
strContent=""
arrLines=[]
for oneDoc in arrDocx :
    list_paras = file_docx_python.Get_Paragraph_List_In_Doc('\\'.join(os.path.realpath(__file__).split('\\')[:-1])+"\\"+oneDoc)
    for para in list_paras:
        if len(para)>100 :
            strContent=para 
            arrLines+=strContent.split("老友记.本章节注释")[0].split("\n")
     
arrCH=[]
arrENG=[]
for line in arrLines:
    if basic_function_python.Check_Chinese_In_String(line):
        arrCH.append(line)
    else :
        arrENG.append(line)

        
str_eng='\n'.join(arrENG)
str_ch = '\n'.join(arrCH)



print([d for d in os.listdir('.')] )


#file_stream_python. Write_Content_To_File(str_eng,"eng.txt")
#file_stream_python. Write_Content_To_File(str_ch,"ch.txt")

#doc_name=file_docx_python.Create_Document("bb.docx")




#file_docx_python.Add_Paragraph_String_To_Document(str_eng,doc_name)
#ile_docx_python.Add_Paragraph_String_To_Document(str_ch,doc_name)


"""
document=file_docx_python.Create_Document()
file_docx_python.Add_Heading_And_Set_Style("asdfasfdasdfadsfasdf",1,document)
file_docx_python.Add_Heading_And_Set_Style("这里不填标题内容",3,document)
file_docx_python.Add_Page_Break(document)
file_docx_python.Add_Table_Use_Rows_And_Cols(2,2,document)
file_docx_python.Add_Page_Break(document)
file_docx_python.Add_Picture_And_Set_Size("usa.jpg",5,document)
file_docx_python.Set_Document_Style(document)

file_docx_python.Save_Document("caowei.docx",document)
"""
  
